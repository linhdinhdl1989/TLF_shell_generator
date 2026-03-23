"""
TLF Shell Generator — FastAPI backend.

Start with:
    uvicorn app:app --reload

Quick smoke-test sequence:
    # 1. Create a study
    curl -s -X POST http://localhost:8000/studies \
        -H "Content-Type: application/json" \
        -d '{"name": "StudyXYZ"}' | jq .

    # 2. Upload a document (multipart)
    curl -s -X POST http://localhost:8000/studies/{study_id}/documents \
        -F "file=@sap.pdf" \
        -F "type=sap" | jq .

    # 3. Get the TLF list (empty until extracted/added)
    curl -s http://localhost:8000/studies/{study_id}/tlf-list | jq .

    # 4. Add a TLF entry manually
    curl -s -X POST http://localhost:8000/studies/{study_id}/tlf-list \
        -H "Content-Type: application/json" \
        -d '{"number":"14.1.1","title":"Demographics","type":"table","section_ref":"SAP 12.1"}' | jq .

    # 5. Create a shell for a TLF
    curl -s -X POST http://localhost:8000/studies/{study_id}/shells \
        -H "Content-Type: application/json" \
        -d '{"tlf_id":"{tlf_id}","title":"Demographics by ARM"}' | jq .

    # 6. Send a chat message
    curl -s -X POST http://localhost:8000/studies/{study_id}/chat \
        -H "Content-Type: application/json" \
        -d '{"prompt":"Add a footnote referencing SAP Section 12.1"}' | jq .

All routes are scoped to a study_id for RLS-readiness.
"""

from __future__ import annotations

import csv
import io
import json
import os
import uuid
import zipfile
from contextlib import asynccontextmanager
from datetime import datetime
from typing import List, Optional

from fastapi import (
    BackgroundTasks,
    Depends,
    FastAPI,
    File,
    Form,
    HTTPException,
    Query,
    UploadFile,
    status,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from sqlalchemy import select, delete, func
from sqlalchemy.ext.asyncio import AsyncSession

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "services"))

from database import get_db, init_db
from models import Action, Document, GlobalRequirement, Message, Shell, Study, TLF, TLFListItem
from schemas import (
    ActionCreate,
    ActionList,
    ActionRead,
    BulkUpdateAnalysisSetRequest,
    AuditEventList,
    AuditEventRead,
    ChatRequest,
    ChatResponse,
    DocumentList,
    DocumentRead,
    DocumentType,
    DocumentUpdate,
    GlobalRequirementBulkUpdate,
    GlobalRequirementCreate,
    GlobalRequirementList,
    GlobalRequirementRead,
    GlobalRequirementUpdate,
    MessageList,
    MessageRead,
    ShellCreate,
    ShellList,
    ShellRead,
    ShellUpdate,
    StudyCreate,
    StudyList,
    StudyRead,
    TLFCreate,
    TLFList,
    TLFListBulkUpdate,
    TLFListItemCreate,
    TLFListItemRead,
    TLFListItemUpdate,
    TLFListResponse,
    TLFRead,
    TLFUpdate,
)


# ---------------------------------------------------------------------------
# App lifecycle
# ---------------------------------------------------------------------------

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Initialise DB tables on startup (dev/SQLite).  Use Alembic in prod."""
    await init_db()
    yield


app = FastAPI(
    title="TLF Shell Generator API",
    description=(
        "Backend for the AI-powered TLF mock shell generator. "
        "All routes are study-scoped for RLS readiness.  "
        "See PRD v1.8 for full feature description."
    ),
    version="0.1.0",
    lifespan=lifespan,
)

# Allow origins from environment or default dev ports
_raw_origins = os.getenv(
    "CORS_ORIGINS",
    "http://localhost:3000,http://localhost:5000,http://localhost:5173"
)
_ALLOWED_ORIGINS: List[str] = [origin.strip() for origin in _raw_origins.split(",") if origin.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=_ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _not_found(entity: str, id_: str) -> HTTPException:
    return HTTPException(status_code=404, detail=f"{entity} '{id_}' not found.")


async def _get_study(study_id: str, db: AsyncSession) -> Study:
    row = await db.get(Study, study_id)
    if row is None:
        raise _not_found("Study", study_id)
    return row


async def _get_tlf(study_id: str, tlf_id: str, db: AsyncSession) -> TLF:
    row = await db.get(TLF, tlf_id)
    if row is None or row.study_id != study_id:
        raise _not_found("TLF", tlf_id)
    return row


async def _get_shell(study_id: str, shell_id: str, db: AsyncSession) -> Shell:
    row = await db.get(Shell, shell_id)
    if row is None or row.study_id != study_id:
        raise _not_found("Shell", shell_id)
    return row


async def _log_action(
    db: AsyncSession,
    *,
    study_id: str,
    action_type: str,
    tlf_id: Optional[str] = None,
    shell_id: Optional[str] = None,
    target: Optional[str] = None,
    before: Optional[dict] = None,
    after: Optional[dict] = None,
    actor: str = "user",
) -> Action:
    """Create an Action (audit trail) row inside the current session."""
    # Pre-compute enriched fields so they're stored and queryable
    entity_type = _derive_entity_type(action_type, target, shell_id, tlf_id)
    summary = _compute_action_summary(action_type, target, after or {}, before or {})
    entry = Action(
        id=str(uuid.uuid4()),
        study_id=study_id,
        tlf_id=tlf_id,
        shell_id=shell_id,
        type=action_type,
        target=target,
        before_state=before,
        after_state=after,
        actor=actor,
        timestamp=datetime.utcnow(),
        summary=summary,
        entity_type=entity_type,
    )
    db.add(entry)
    return entry


def _derive_entity_type(
    action_type: str,
    target: Optional[str],
    shell_id: Optional[str],
    tlf_id: Optional[str],
) -> str:
    """Compute entity_type for an audit event from available context."""
    if action_type in (
        "shell_created", "update_title", "update_subtitle", "update_population",
        "update_column", "update_row", "update_footnote", "update_status",
        "approve_shell", "reject_shell", "ai_reviewer_correction",
    ) and shell_id:
        return "shell"
    if shell_id:
        return "shell"
    if action_type in ("tlf_list_saved", "tlf_list_extracted", "tlf_list_approved"):
        return "tlf_list"
    if action_type in ("add_row", "delete_row", "update_row", "reorder_rows") or target == "tlf_list":
        return "tlf_list"
    if action_type == "global_reqs_saved" or target == "global_requirements":
        return "global_requirements"
    if action_type == "document_uploaded":
        return "document"
    if action_type == "chat_sent":
        return "chat"
    if tlf_id:
        return "tlf_list"
    return "study"


def _compute_action_summary(
    action_type: str,
    target: Optional[str],
    after: dict,
    before: dict,
) -> str:
    """Generate a concise, human-readable summary for display in the audit trail."""
    t = target or ""
    if action_type == "shell_created":
        title = after.get("title", "New Shell")
        return f'Shell created: "{title}"'
    if action_type in ("tlf_list_saved", "reorder_rows"):
        tlfs = after.get("tlfs", [])
        count = len(tlfs) if isinstance(tlfs, list) else after.get("count", "?")
        return f"TLF list saved ({count} entries)"
    if action_type == "tlf_list_extracted":
        return f"TLF list extracted by AI ({after.get('extracted_count', '?')} entries)"
    if action_type == "tlf_list_approved":
        return f"TLF list approved ({after.get('count', '?')} entries)"
    if action_type == "global_reqs_saved":
        return f"Global requirements saved ({after.get('count', '?')} sections)"
    if action_type == "document_uploaded":
        name = after.get("name", "")
        doc_type = after.get("type", "")
        return f"Document uploaded: {name}" + (f" ({doc_type})" if doc_type else "")
    if action_type == "chat_sent":
        prompt = after.get("prompt", t)
        short = str(prompt)[:60]
        return f'Chat sent: "{short}"' + ("…" if len(str(prompt)) > 60 else "")
    if action_type == "add_row":
        title = after.get("title", after.get("number", ""))
        return f"TLF entry added: {title}" if title else "TLF entry added"
    if action_type == "delete_row":
        title = before.get("title", before.get("number", ""))
        return f"TLF entry deleted: {title}" if title else "TLF entry deleted"
    if action_type == "update_row":
        if t.startswith("tlf:"):
            return f"TLF entry updated: {t[4:]}"
        return "Shell rows updated"
    if action_type == "update_title":
        val = str(after.get("title", ""))[:50]
        return f'Shell title updated: "{val}"'
    if action_type == "update_subtitle":
        return "Shell subtitle updated"
    if action_type == "update_population":
        val = after.get("population", "")
        return f"Shell population: {val}" if val else "Shell population updated"
    if action_type == "update_column":
        return "Shell columns updated"
    if action_type == "update_footnote":
        return "Shell footnotes updated"
    if action_type == "update_status":
        new_status = after.get("status", "")
        if t == "tlf_list":
            return f"TLF list status changed to {new_status}" if new_status else "TLF list status updated"
        return f"Shell status changed to {new_status}" if new_status else "Shell status updated"
    if action_type == "approve_shell":
        return "Shell approved"
    if action_type == "reject_shell":
        return "Shell deleted"
    if action_type == "ai_suggestion":
        if t == "tlf_list":
            return f"TLF list extracted by AI ({after.get('extracted_count', '?')} entries)"
        if t == "shell":
            return f'Shell created by AI: "{after.get("title", "New Shell")}"'
        return "AI suggestion applied"
    if action_type == "ai_variable_flagged":
        return "AI flagged variable"
    if action_type == "ai_reviewer_correction":
        return "AI reviewer correction"
    if action_type == "add_column":
        return "Column added"
    if action_type == "delete_column":
        return "Column removed"
    # Generic fallback
    return action_type.replace("_", " ").title()


def _action_to_event(action: Action) -> AuditEventRead:
    """Convert a stored Action row into an enriched AuditEventRead."""
    after = action.after_state or {}
    before = action.before_state or {}

    entity_type = action.entity_type or _derive_entity_type(
        action.type, action.target, action.shell_id, action.tlf_id
    )
    summary = action.summary or _compute_action_summary(
        action.type, action.target, after, before
    )

    # Determine entity_id
    entity_id: Optional[str] = None
    if entity_type == "shell":
        entity_id = action.shell_id
    elif entity_type in ("tlf_list",):
        entity_id = action.tlf_id
    elif entity_type == "document":
        entity_id = after.get("doc_id")

    source = "ai" if action.actor == "ai" else "user"

    return AuditEventRead(
        id=action.id,
        study_id=action.study_id,
        shell_id=action.shell_id,
        tlf_id=action.tlf_id,
        entity_type=entity_type,
        entity_id=entity_id,
        action=action.type,
        summary=summary,
        details=after if after else (before if before else None),
        actor=action.actor,
        source=source,
        created_at=action.timestamp,
    )


# ---------------------------------------------------------------------------
# Background task: document parse stub  (PRD §4.7)
# ---------------------------------------------------------------------------

async def _parse_document_stub(doc_id: str) -> None:
    """
    Stub for the parse-once pipeline (PRD §4.7).

    Real implementation would:
      1. Retrieve raw bytes from temporary storage.
      2. Extract text (pdfplumber / python-docx / openpyxl).
      3. Fallback OCR for scanned PDFs.
      4. Split into chunks with section_ref tags.
      5. Embed chunks (pgvector / external API).
      6. Store parsed_content + chunks_meta on Document row.
      7. Delete raw file.
      8. Set status = "ready" (or "error" on failure).

    The stub simply marks the document as "ready" after a simulated delay.
    Replace with a Celery/ARQ/BackgroundTasks task in production.
    """
    import asyncio

    # Simulate I/O work (parsing, embedding)
    await asyncio.sleep(0.5)

    # Re-open a fresh session — background tasks run outside the request session
    from database import AsyncSessionLocal

    async with AsyncSessionLocal() as session:
        doc = await session.get(Document, doc_id)
        if doc is None:
            return
        if doc.status == "processing":
            doc.parsed_content = (
                f"[STUB] Parsed content for document '{doc.name}' "
                "(replace with real extraction pipeline)"
            )
            doc.chunks_meta = [
                {
                    "chunk_id": str(uuid.uuid4()),
                    "section_ref": "Section 1",
                    "text": f"[STUB] Chunk 1 from {doc.name}",
                    "embedding_id": None,
                }
            ]
            doc.status = "ready"
            doc.updated_at = datetime.utcnow()
            await session.commit()


# ---------------------------------------------------------------------------
# Root
# ---------------------------------------------------------------------------

@app.get("/", tags=["Health"])
async def root():
    """Health check / API root."""
    return {
        "service": "TLF Shell Generator API",
        "version": "0.1.0",
        "status": "ok",
        "docs": "/docs",
    }


@app.get("/health", tags=["Health"])
async def health():
    return {"status": "ok"}


# ===========================================================================
# Studies
# ===========================================================================

@app.post(
    "/studies",
    response_model=StudyRead,
    status_code=status.HTTP_201_CREATED,
    tags=["Studies"],
    summary="Create a new study",
)
async def create_study(
    body: StudyCreate,
    db: AsyncSession = Depends(get_db),
):
    """
    POST /studies  →  {id, name, created_at, updated_at}

    Example:
        curl -X POST /studies -d '{"name": "StudyXYZ"}'
    """
    study = Study(id=str(uuid.uuid4()), name=body.name)
    db.add(study)
    await db.flush()
    await db.refresh(study)
    return study


@app.get(
    "/studies",
    response_model=StudyList,
    tags=["Studies"],
    summary="List all studies",
)
async def list_studies(
    skip: int = Query(0, ge=0),
    limit: int = Query(50, ge=1, le=200),
    db: AsyncSession = Depends(get_db),
):
    total_result = await db.execute(select(func.count(Study.id)))
    total = total_result.scalar_one()

    result = await db.execute(
        select(Study).order_by(Study.created_at.desc()).offset(skip).limit(limit)
    )
    rows = result.scalars().all()
    return StudyList(studies=list(rows), total=total)


@app.get(
    "/studies/{study_id}",
    response_model=StudyRead,
    tags=["Studies"],
    summary="Get a study by ID",
)
async def get_study(study_id: str, db: AsyncSession = Depends(get_db)):
    return await _get_study(study_id, db)


# ===========================================================================
# Documents
# ===========================================================================

@app.post(
    "/studies/{study_id}/documents",
    response_model=DocumentRead,
    status_code=status.HTTP_201_CREATED,
    tags=["Documents"],
    summary="Upload a document (multipart)",
)
async def upload_document(
    study_id: str,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    type: DocumentType = Form(...),
    label: Optional[str] = Form(None, description="Human label for 'other' type"),
    db: AsyncSession = Depends(get_db),
):
    """
    POST /studies/{study_id}/documents

    Multipart form fields:
      - file   : the binary file
      - type   : sap | protocol | tlf_library | study_tlf_list | other
      - label  : (optional) free-text label for 'other' documents

    Returns the document record with status='processing'.
    A background task moves it to 'ready' after parsing.

    Example:
        curl -X POST /studies/{id}/documents \\
            -F "file=@sap.pdf" \\
            -F "type=sap"
    """
    await _get_study(study_id, db)

    file_bytes = await file.read()
    doc = Document(
        id=str(uuid.uuid4()),
        study_id=study_id,
        name=file.filename or "untitled",
        original_filename=file.filename,
        type=type.value,
        label=label,
        status="processing",
        file_size=len(file_bytes),
        content_type=file.content_type,
    )
    db.add(doc)
    await db.flush()
    await db.refresh(doc)

    await _log_action(
        db,
        study_id=study_id,
        action_type="document_uploaded",
        target=f"document:{doc.id}",
        after={"name": doc.name, "type": doc.type, "doc_id": doc.id, "file_size": doc.file_size},
    )

    # Fire-and-forget parse pipeline (stub)
    background_tasks.add_task(_parse_document_stub, doc.id)

    return doc


@app.get(
    "/studies/{study_id}/documents",
    response_model=DocumentList,
    tags=["Documents"],
    summary="List documents for a study",
)
async def list_documents(
    study_id: str,
    db: AsyncSession = Depends(get_db),
):
    await _get_study(study_id, db)
    result = await db.execute(
        select(Document)
        .where(Document.study_id == study_id)
        .order_by(Document.created_at)
    )
    rows = result.scalars().all()
    return DocumentList(documents=list(rows), total=len(rows))


@app.get(
    "/studies/{study_id}/documents/{doc_id}",
    response_model=DocumentRead,
    tags=["Documents"],
    summary="Get a single document",
)
async def get_document(
    study_id: str,
    doc_id: str,
    db: AsyncSession = Depends(get_db),
):
    await _get_study(study_id, db)
    doc = await db.get(Document, doc_id)
    if doc is None or doc.study_id != study_id:
        raise _not_found("Document", doc_id)
    return doc


@app.put(
    "/studies/{study_id}/documents/{doc_id}",
    response_model=DocumentRead,
    tags=["Documents"],
    summary="Update document metadata (name, label, type)",
)
async def update_document(
    study_id: str,
    doc_id: str,
    body: DocumentUpdate,
    db: AsyncSession = Depends(get_db),
):
    await _get_study(study_id, db)
    doc = await db.get(Document, doc_id)
    if doc is None or doc.study_id != study_id:
        raise _not_found("Document", doc_id)

    if body.name is not None:
        doc.name = body.name
    if body.label is not None:
        doc.label = body.label
    if body.type is not None:
        doc.type = body.type.value
    doc.updated_at = datetime.utcnow()

    await db.flush()
    await db.refresh(doc)
    return doc


# ===========================================================================
# TLF List
# ===========================================================================

@app.get(
    "/studies/{study_id}/tlf-list",
    response_model=TLFList,
    tags=["TLF List"],
    summary="Get the TLF list for a study",
)
async def get_tlf_list(
    study_id: str,
    db: AsyncSession = Depends(get_db),
):
    """
    GET /studies/{study_id}/tlf-list

    Returns an ordered array of TLF entries.
    If no TLF list has been uploaded, this returns an empty list.
    (Populate it via the POST endpoint or the bulk PUT endpoint.)

    Example:
        curl http://localhost:8000/studies/{study_id}/tlf-list
    """
    await _get_study(study_id, db)
    result = await db.execute(
        select(TLF)
        .where(TLF.study_id == study_id)
        .order_by(TLF.order_index, TLF.number)
    )
    rows = result.scalars().all()
    return TLFList(tlfs=list(rows), total=len(rows))


@app.post(
    "/studies/{study_id}/tlf-list",
    response_model=TLFRead,
    status_code=status.HTTP_201_CREATED,
    tags=["TLF List"],
    summary="Add a single TLF entry",
)
async def add_tlf(
    study_id: str,
    body: TLFCreate,
    db: AsyncSession = Depends(get_db),
):
    await _get_study(study_id, db)
    tlf = TLF(
        id=str(uuid.uuid4()),
        study_id=study_id,
        **body.model_dump(),
    )
    db.add(tlf)
    await _log_action(
        db,
        study_id=study_id,
        action_type="add_row",
        tlf_id=tlf.id,
        target=f"tlf:{tlf.number}",
        after={"number": tlf.number, "title": tlf.title},
    )
    await db.flush()
    await db.refresh(tlf)
    return tlf


@app.put(
    "/studies/{study_id}/tlf-list",
    response_model=TLFList,
    tags=["TLF List"],
    summary="Bulk-replace the TLF list (editable table save)",
)
async def bulk_replace_tlf_list(
    study_id: str,
    body: TLFListBulkUpdate,
    db: AsyncSession = Depends(get_db),
):
    """
    PUT /studies/{study_id}/tlf-list

    Replaces the entire TLF list.  Existing TLFs for this study are deleted
    and re-created from the payload.  Use this for the 'Save TLF List' action
    in the tabbed editor.

    All order_index values should be provided by the client to preserve row order.
    """
    await _get_study(study_id, db)

    # Snapshot before state for audit
    existing_result = await db.execute(
        select(TLF).where(TLF.study_id == study_id).order_by(TLF.order_index)
    )
    before_tlfs = [
        {"number": t.number, "title": t.title, "status": t.status}
        for t in existing_result.scalars().all()
    ]

    await db.execute(delete(TLF).where(TLF.study_id == study_id))

    new_tlfs: List[TLF] = []
    for item in body.tlfs:
        tlf = TLF(id=str(uuid.uuid4()), study_id=study_id, **item.model_dump())
        db.add(tlf)
        new_tlfs.append(tlf)

    after_tlfs = [{"number": t.number, "title": t.title} for t in new_tlfs]

    await _log_action(
        db,
        study_id=study_id,
        action_type="tlf_list_saved",
        target="tlf_list",
        before={"tlfs": before_tlfs},
        after={"tlfs": after_tlfs},
    )

    await db.flush()
    for tlf in new_tlfs:
        await db.refresh(tlf)

    return TLFList(tlfs=new_tlfs, total=len(new_tlfs))


@app.put(
    "/studies/{study_id}/tlf-list/{tlf_id}",
    response_model=TLFRead,
    tags=["TLF List"],
    summary="Update a single TLF entry",
)
async def update_tlf(
    study_id: str,
    tlf_id: str,
    body: TLFUpdate,
    db: AsyncSession = Depends(get_db),
):
    tlf = await _get_tlf(study_id, tlf_id, db)

    before = {"number": tlf.number, "title": tlf.title, "status": tlf.status}

    for field, value in body.model_dump(exclude_none=True).items():
        if hasattr(tlf, field) and value is not None:
            if hasattr(value, "value"):
                setattr(tlf, field, value.value)
            else:
                setattr(tlf, field, value)

    tlf.updated_at = datetime.utcnow()
    after = {"number": tlf.number, "title": tlf.title, "status": tlf.status}

    await _log_action(
        db,
        study_id=study_id,
        action_type="update_row",
        tlf_id=tlf_id,
        target=f"tlf:{tlf.number}",
        before=before,
        after=after,
    )

    await db.flush()
    await db.refresh(tlf)
    return tlf


@app.delete(
    "/studies/{study_id}/tlf-list/{tlf_id}",
    status_code=status.HTTP_204_NO_CONTENT,
    tags=["TLF List"],
    summary="Delete a TLF entry",
)
async def delete_tlf(
    study_id: str,
    tlf_id: str,
    db: AsyncSession = Depends(get_db),
):
    tlf = await _get_tlf(study_id, tlf_id, db)

    await _log_action(
        db,
        study_id=study_id,
        action_type="delete_row",
        tlf_id=tlf_id,
        target=f"tlf:{tlf.number}",
        before={"number": tlf.number, "title": tlf.title},
    )

    await db.delete(tlf)


@app.post(
    "/studies/{study_id}/tlf-list/extract",
    response_model=TLFList,
    status_code=status.HTTP_201_CREATED,
    tags=["TLF List"],
    summary="AI-extract candidate TLF rows from SAP document (or stub)",
)
async def extract_tlf_list(
    study_id: str,
    db: AsyncSession = Depends(get_db),
):
    """
    POST /studies/{study_id}/tlf-list/extract

    Inspects the study's parsed SAP document and returns a candidate TLF list.
    If no SAP document is available, a realistic clinical stub list is returned.

    The extracted rows are persisted as proposed TLF entries, replacing any
    existing list.  Users can review and edit before approving.
    """
    await _get_study(study_id, db)

    # Check for a ready SAP document to build context
    sap_result = await db.execute(
        select(Document)
        .where(Document.study_id == study_id)
        .where(Document.type == "sap")
        .where(Document.status == "ready")
        .order_by(Document.created_at.desc())
    )
    sap_doc = sap_result.scalars().first()
    source = "sap_stub" if sap_doc else "clinical_stub"

    # Realistic clinical stub — mirrors a typical Phase 3 CTD Section 14 TLF list.
    # Replace with real NLP/LLM extraction in production.
    stub_rows = [
        {"number": "14.1.1", "title": "Demographics and Baseline Characteristics", "section_ref": "demographics"},
        {"number": "14.1.2", "title": "Medical History Summary", "section_ref": "demographics"},
        {"number": "14.1.3", "title": "Prior and Concomitant Medications", "section_ref": "demographics"},
        {"number": "14.2.1", "title": "Primary Efficacy Endpoint – Change from Baseline", "section_ref": "efficacy"},
        {"number": "14.2.2", "title": "Secondary Efficacy – Responder Analysis (≥30% Improvement)", "section_ref": "efficacy"},
        {"number": "14.2.3", "title": "Time to First Response – Kaplan-Meier Analysis", "section_ref": "efficacy"},
        {"number": "14.3.1.1", "title": "Adverse Events – Overview", "section_ref": "safety"},
        {"number": "14.3.1.2", "title": "Treatment-Emergent Adverse Events by SOC and PT", "section_ref": "safety"},
        {"number": "14.3.2.1", "title": "Serious Adverse Events", "section_ref": "safety"},
        {"number": "14.3.2.2", "title": "Adverse Events Leading to Discontinuation", "section_ref": "safety"},
        {"number": "14.3.3.1", "title": "Clinical Laboratory Parameters – Summary Statistics", "section_ref": "safety"},
        {"number": "14.3.3.2", "title": "Clinically Notable Laboratory Values", "section_ref": "safety"},
    ]

    # Delete current TLF list and replace with extracted candidate rows
    await db.execute(delete(TLF).where(TLF.study_id == study_id))

    new_tlfs: List[TLF] = []
    for i, row in enumerate(stub_rows):
        tlf = TLF(
            id=str(uuid.uuid4()),
            study_id=study_id,
            number=row["number"],
            title=row["title"],
            type="table",
            section_ref=row["section_ref"],
            status="proposed",
            order_index=i,
        )
        db.add(tlf)
        new_tlfs.append(tlf)

    await _log_action(
        db,
        study_id=study_id,
        action_type="tlf_list_extracted",
        target="tlf_list",
        after={"extracted_count": len(new_tlfs), "source": source},
        actor="ai",
    )

    await db.flush()
    for tlf in new_tlfs:
        await db.refresh(tlf)

    return TLFList(tlfs=new_tlfs, total=len(new_tlfs))


@app.post(
    "/studies/{study_id}/tlf-list/approve",
    response_model=TLFList,
    tags=["TLF List"],
    summary="Approve the entire TLF list for a study",
)
async def approve_tlf_list(
    study_id: str,
    db: AsyncSession = Depends(get_db),
):
    """
    POST /studies/{study_id}/tlf-list/approve

    Marks every TLF entry for this study as 'approved'.
    Returns the updated list.
    """
    await _get_study(study_id, db)

    result = await db.execute(
        select(TLF)
        .where(TLF.study_id == study_id)
        .order_by(TLF.order_index, TLF.number)
    )
    tlfs = result.scalars().all()

    for tlf in tlfs:
        tlf.status = "approved"
        tlf.updated_at = datetime.utcnow()

    await _log_action(
        db,
        study_id=study_id,
        action_type="tlf_list_approved",
        target="tlf_list",
        after={"status": "approved", "count": len(tlfs)},
    )

    await db.flush()
    for tlf in tlfs:
        await db.refresh(tlf)

    return TLFList(tlfs=list(tlfs), total=len(tlfs))


# ===========================================================================
# TLFListItem endpoints (new normalized TLF list with provenance)
# ===========================================================================

def _get_stub_extraction_candidates():
    """Return stub clinical TLF candidates when no SAP content is available."""
    return [
        {"number": "14.1.1", "raw_title": "Summary of Demographic and Baseline Characteristics", "output_type": "table", "section": "demographics",
         "extraction_evidence": {"subtitle_context": "By Treatment Group", "analysis_set_context": "Safety Population", "title_text_source": "stub"}},
        {"number": "14.1.2", "raw_title": "Medical History Summary", "output_type": "table", "section": "demographics",
         "extraction_evidence": {"subtitle_context": "By Treatment Group", "analysis_set_context": "Safety Population", "title_text_source": "stub"}},
        {"number": "14.1.3", "raw_title": "Prior and Concomitant Medications", "output_type": "table", "section": "demographics",
         "extraction_evidence": {"subtitle_context": None, "analysis_set_context": "Safety Population", "title_text_source": "stub"}},
        {"number": "14.2.1", "raw_title": "Primary Efficacy Endpoint - Change from Baseline", "output_type": "table", "section": "efficacy",
         "extraction_evidence": {"subtitle_context": "By Treatment Group", "analysis_set_context": "Full Analysis Set", "title_text_source": "stub"}},
        {"number": "14.2.2", "raw_title": "Responder Analysis (>=30% Improvement)", "output_type": "table", "section": "efficacy",
         "extraction_evidence": {"subtitle_context": "By Treatment Group", "analysis_set_context": "Full Analysis Set", "title_text_source": "stub"}},
        {"number": "14.2.3", "raw_title": "Time to First Response", "output_type": "figure", "section": "efficacy",
         "extraction_evidence": {"subtitle_context": "Kaplan-Meier Analysis", "analysis_set_context": "Full Analysis Set", "title_text_source": "stub"}},
        {"number": "14.3.1.1", "raw_title": "Adverse Events - Overview", "output_type": "table", "section": "safety",
         "extraction_evidence": {"subtitle_context": None, "analysis_set_context": "Safety Population", "title_text_source": "stub"}},
        {"number": "14.3.1.2", "raw_title": "Treatment-Emergent Adverse Events", "output_type": "table", "section": "safety",
         "extraction_evidence": {"subtitle_context": "By SOC and PT", "analysis_set_context": "Safety Population", "title_text_source": "stub"}},
        {"number": "14.3.2.1", "raw_title": "Serious Adverse Events", "output_type": "table", "section": "safety",
         "extraction_evidence": {"subtitle_context": None, "analysis_set_context": "Safety Population", "title_text_source": "stub"}},
        {"number": "14.3.2.2", "raw_title": "Adverse Events Leading to Discontinuation", "output_type": "table", "section": "safety",
         "extraction_evidence": {"subtitle_context": None, "analysis_set_context": "Safety Population", "title_text_source": "stub"}},
        {"number": "14.3.3.1", "raw_title": "Clinical Laboratory Parameters - Summary Statistics", "output_type": "table", "section": "safety",
         "extraction_evidence": {"subtitle_context": "By Treatment Group", "analysis_set_context": "Safety Population", "title_text_source": "stub"}},
        {"number": "14.3.3.2", "raw_title": "Clinically Notable Laboratory Values", "output_type": "table", "section": "safety",
         "extraction_evidence": {"subtitle_context": None, "analysis_set_context": "Safety Population", "title_text_source": "stub"}},
    ]


@app.get("/studies/{study_id}/tlf-list/items", response_model=TLFListResponse, tags=["TLF List Items"])
async def get_tlf_list_items(study_id: str, db: AsyncSession = Depends(get_db)):
    await _get_study(study_id, db)
    result = await db.execute(
        select(TLFListItem).where(TLFListItem.study_id == study_id).order_by(TLFListItem.order_index, TLFListItem.number)
    )
    items = result.scalars().all()
    approved_count = sum(1 for i in items if i.approved)
    return TLFListResponse(items=list(items), total=len(items), approved_count=approved_count)


@app.post("/studies/{study_id}/tlf-list/extract-items", response_model=TLFListResponse, status_code=201, tags=["TLF List Items"])
async def extract_tlf_list_items(study_id: str, db: AsyncSession = Depends(get_db)):
    """Extract candidate TLF list from SAP parsed content, normalize, and store."""
    from services.tlf_extraction_service import extract_tlf_candidates_from_sap
    from services.tlf_title_normalizer import normalize_tlf_title

    await _get_study(study_id, db)

    # Get parsed SAP content
    sap_result = await db.execute(
        select(Document).where(Document.study_id == study_id)
        .where(Document.type == "sap").where(Document.status == "ready")
        .order_by(Document.created_at.desc())
    )
    sap_doc = sap_result.scalars().first()

    if not sap_doc:
        candidates = _get_stub_extraction_candidates()
        source = "clinical_stub"
    else:
        parsed_content = getattr(sap_doc, 'parsed_content', '') or ''
        if not parsed_content:
            candidates = _get_stub_extraction_candidates()
            source = "sap_stub"
        else:
            candidates = extract_tlf_candidates_from_sap(parsed_content)
            source = "extracted"
        if not candidates:
            candidates = _get_stub_extraction_candidates()
            source = "clinical_stub"

    # Delete existing extracted items for this study
    await db.execute(delete(TLFListItem).where(TLFListItem.study_id == study_id))

    new_items = []
    for i, cand in enumerate(candidates):
        normalized = normalize_tlf_title(
            raw_title=cand.get("raw_title"),
            extraction_evidence=cand.get("extraction_evidence"),
        )
        item = TLFListItem(
            id=str(uuid.uuid4()),
            study_id=study_id,
            number=cand["number"],
            output_type=cand.get("output_type", "table"),
            section=cand.get("section", "other"),
            raw_title=cand.get("raw_title"),
            source=source,
            extraction_notes=str(cand.get("extraction_evidence", {})),
            title=normalized["title"],
            subtitle=normalized["subtitle"],
            analysis_set=normalized["analysis_set"],
            composed_title=normalized["composed_title"],
            title_source=normalized["title_source"],
            subtitle_source=normalized["subtitle_source"],
            analysis_set_source=normalized["analysis_set_source"],
            parsing_confidence=normalized["parsing_confidence"],
            status="pending",
            approved=False,
            order_index=i,
        )
        db.add(item)
        new_items.append(item)

    await _log_action(db, study_id=study_id, action_type="ai_suggestion", target="tlf_list_items",
                      after={"extracted_count": len(new_items), "source": source}, actor="ai")
    await db.flush()
    for item in new_items:
        await db.refresh(item)

    approved_count = 0
    return TLFListResponse(items=new_items, total=len(new_items), approved_count=approved_count)


@app.post("/studies/{study_id}/tlf-list/upload-items", response_model=TLFListResponse, status_code=201, tags=["TLF List Items"])
async def upload_tlf_list_items(
    study_id: str,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    """Upload CSV or Excel TLF list. Normalizes each row and stores."""
    from services.tlf_title_normalizer import normalize_uploaded_row
    from services.tlf_extraction_service import _infer_section
    import io

    await _get_study(study_id, db)

    content = await file.read()
    filename = file.filename or ""
    rows = []
    warnings = []

    try:
        if filename.endswith(".csv"):
            import csv
            reader = csv.DictReader(io.StringIO(content.decode("utf-8-sig")))
            rows = list(reader)
        elif filename.endswith((".xlsx", ".xls")):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(io.BytesIO(content))
                ws = wb.active
                headers = [str(c.value or "").strip().lower().replace(" ", "_") for c in next(ws.iter_rows(max_row=1))]
                for row in ws.iter_rows(min_row=2, values_only=True):
                    rows.append({headers[j]: (str(v) if v is not None else "") for j, v in enumerate(row)})
            except ImportError:
                raise HTTPException(status_code=400, detail="openpyxl not installed. Upload CSV instead.")
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Use .csv or .xlsx")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse file: {e}")

    # Delete existing items
    await db.execute(delete(TLFListItem).where(TLFListItem.study_id == study_id))

    new_items = []
    for i, row in enumerate(rows):
        if not any(row.values()):
            continue

        # Normalize column names
        normalized_row = {}
        for k, v in row.items():
            clean_k = k.strip().lower().replace(" ", "_").replace("-", "_")
            normalized_row[clean_k] = str(v).strip() if v else ""

        number = normalized_row.get("number") or normalized_row.get("#") or normalized_row.get("tlf_number") or ""
        if not number:
            warnings.append(f"Row {i+2}: missing TLF number, skipped")
            continue

        norm = normalize_uploaded_row(normalized_row)

        # Infer section and output_type if not provided
        section = normalized_row.get("section") or _infer_section(number, norm.get("title") or "", "")
        output_type = normalized_row.get("output_type") or normalized_row.get("type") or "table"

        item = TLFListItem(
            id=str(uuid.uuid4()),
            study_id=study_id,
            number=number,
            output_type=output_type,
            section=section,
            raw_title=normalized_row.get("raw_title") or normalized_row.get("complete_title"),
            source="uploaded",
            title=norm["title"] or number,
            subtitle=norm["subtitle"],
            analysis_set=norm["analysis_set"],
            composed_title=norm["composed_title"] or norm["title"] or number,
            title_source=norm["title_source"],
            subtitle_source=norm["subtitle_source"],
            analysis_set_source=norm["analysis_set_source"],
            parsing_confidence=norm["parsing_confidence"],
            status="pending",
            approved=False,
            order_index=i,
        )
        db.add(item)
        new_items.append(item)

    await db.flush()
    for item in new_items:
        await db.refresh(item)

    approved_count = 0
    return TLFListResponse(items=new_items, total=len(new_items), approved_count=approved_count)


@app.post("/studies/{study_id}/tlf-list/items", response_model=TLFListItemRead, status_code=201, tags=["TLF List Items"])
async def create_tlf_list_item(study_id: str, body: TLFListItemCreate, db: AsyncSession = Depends(get_db)):
    from services.tlf_title_normalizer import _build_composed_title
    await _get_study(study_id, db)

    # Recompute composed_title
    composed = _build_composed_title(body.title, body.subtitle, body.analysis_set)

    item = TLFListItem(
        id=str(uuid.uuid4()),
        study_id=study_id,
        number=body.number,
        output_type=body.output_type or "table",
        section=body.section or "other",
        raw_title=body.raw_title,
        source="user",
        title=body.title,
        subtitle=body.subtitle,
        analysis_set=body.analysis_set,
        composed_title=composed,
        title_source="user_edit",
        subtitle_source="user_edit" if body.subtitle else None,
        analysis_set_source="user_edit" if body.analysis_set else None,
        parsing_confidence="high" if body.analysis_set else "medium",
        status="pending",
        approved=False,
        order_index=body.order_index,
    )
    db.add(item)
    await db.flush()
    await db.refresh(item)
    return item


@app.put("/studies/{study_id}/tlf-list/items/{item_id}", response_model=TLFListItemRead, tags=["TLF List Items"])
async def update_tlf_list_item(study_id: str, item_id: str, body: TLFListItemUpdate, db: AsyncSession = Depends(get_db)):
    from services.tlf_title_normalizer import _build_composed_title
    await _get_study(study_id, db)
    item = await db.get(TLFListItem, item_id)
    if item is None or item.study_id != study_id:
        raise HTTPException(status_code=404, detail=f"TLFListItem '{item_id}' not found.")

    updates = body.model_dump(exclude_none=True)
    for field, value in updates.items():
        setattr(item, field, value)
        # Track user edits in provenance
        if field == "title":
            item.title_source = "user_edit"
        elif field == "subtitle":
            item.subtitle_source = "user_edit"
        elif field == "analysis_set":
            item.analysis_set_source = "user_edit"

    # Recompute composed_title
    item.composed_title = _build_composed_title(item.title, item.subtitle, item.analysis_set)

    # Update confidence when user explicitly fills in analysis_set
    if "analysis_set" in updates and updates["analysis_set"]:
        item.parsing_confidence = "high"

    item.updated_at = datetime.utcnow()
    await db.flush()
    await db.refresh(item)
    return item


@app.post("/studies/{study_id}/tlf-list/items/{item_id}/approve", response_model=TLFListItemRead, tags=["TLF List Items"])
async def approve_tlf_list_item(study_id: str, item_id: str, db: AsyncSession = Depends(get_db)):
    await _get_study(study_id, db)
    item = await db.get(TLFListItem, item_id)
    if item is None or item.study_id != study_id:
        raise HTTPException(status_code=404, detail=f"TLFListItem '{item_id}' not found.")
    item.approved = True
    item.status = "approved"
    item.updated_at = datetime.utcnow()
    await db.flush()
    await db.refresh(item)
    return item


@app.post("/studies/{study_id}/tlf-list/approve-all-items", response_model=TLFListResponse, tags=["TLF List Items"])
async def approve_all_tlf_list_items(study_id: str, db: AsyncSession = Depends(get_db)):
    await _get_study(study_id, db)
    result = await db.execute(select(TLFListItem).where(TLFListItem.study_id == study_id))
    items = result.scalars().all()
    for item in items:
        item.approved = True
        item.status = "approved"
        item.updated_at = datetime.utcnow()
    await db.flush()
    for item in items:
        await db.refresh(item)
    approved_count = len(items)
    return TLFListResponse(items=list(items), total=len(items), approved_count=approved_count)


@app.delete("/studies/{study_id}/tlf-list/items/{item_id}", status_code=204, tags=["TLF List Items"])
async def delete_tlf_list_item(study_id: str, item_id: str, db: AsyncSession = Depends(get_db)):
    await _get_study(study_id, db)
    item = await db.get(TLFListItem, item_id)
    if item is None or item.study_id != study_id:
        raise HTTPException(status_code=404, detail=f"TLFListItem '{item_id}' not found.")
    await db.delete(item)


@app.post("/studies/{study_id}/tlf-list/bulk-update-analysis-set", tags=["TLF List Items"])
async def bulk_update_analysis_set(study_id: str, body: BulkUpdateAnalysisSetRequest, db: AsyncSession = Depends(get_db)):
    from services.tlf_title_normalizer import _build_composed_title
    await _get_study(study_id, db)
    result = await db.execute(
        select(TLFListItem).where(TLFListItem.study_id == study_id).where(TLFListItem.id.in_(body.item_ids))
    )
    items = result.scalars().all()
    for item in items:
        item.analysis_set = body.analysis_set
        item.analysis_set_source = "user_edit"
        item.composed_title = _build_composed_title(item.title, item.subtitle, item.analysis_set)
        if item.parsing_confidence == "low":
            item.parsing_confidence = "medium"
        item.updated_at = datetime.utcnow()
    await db.flush()
    for item in items:
        await db.refresh(item)
    approved_count = sum(1 for i in items if i.approved)
    return {"updated_count": len(items), "items": [TLFListItemRead.model_validate(i) for i in items]}


# ===========================================================================
# Global Requirements  (shell structure per section type, PRD §4.3)
# ===========================================================================

@app.get(
    "/studies/{study_id}/global-requirements",
    response_model=GlobalRequirementList,
    tags=["Global Requirements"],
    summary="Get global shell structure requirements for a study",
)
async def get_global_requirements(
    study_id: str,
    db: AsyncSession = Depends(get_db),
):
    """
    GET /studies/{study_id}/global-requirements

    Returns per-section shell structure configs (number pattern, title template,
    column headers).  Users configure these before per-shell generation (PRD §4.3).
    """
    await _get_study(study_id, db)
    result = await db.execute(
        select(GlobalRequirement)
        .where(GlobalRequirement.study_id == study_id)
        .order_by(GlobalRequirement.section_type)
    )
    rows = result.scalars().all()
    return GlobalRequirementList(requirements=list(rows), total=len(rows))


@app.post(
    "/studies/{study_id}/global-requirements",
    response_model=GlobalRequirementRead,
    status_code=status.HTTP_201_CREATED,
    tags=["Global Requirements"],
    summary="Create a global requirement for a section type",
)
async def create_global_requirement(
    study_id: str,
    body: GlobalRequirementCreate,
    db: AsyncSession = Depends(get_db),
):
    await _get_study(study_id, db)
    req = GlobalRequirement(
        id=str(uuid.uuid4()),
        study_id=study_id,
        **body.model_dump(),
    )
    db.add(req)
    await db.flush()
    await db.refresh(req)
    return req


@app.put(
    "/studies/{study_id}/global-requirements/{req_id}",
    response_model=GlobalRequirementRead,
    tags=["Global Requirements"],
    summary="Update a global requirement",
)
async def update_global_requirement(
    study_id: str,
    req_id: str,
    body: GlobalRequirementUpdate,
    db: AsyncSession = Depends(get_db),
):
    await _get_study(study_id, db)
    req = await db.get(GlobalRequirement, req_id)
    if req is None or req.study_id != study_id:
        raise _not_found("GlobalRequirement", req_id)

    for field, value in body.model_dump(exclude_none=True).items():
        setattr(req, field, value)
    req.updated_at = datetime.utcnow()

    await db.flush()
    await db.refresh(req)
    return req


@app.delete(
    "/studies/{study_id}/global-requirements/{req_id}",
    status_code=status.HTTP_204_NO_CONTENT,
    tags=["Global Requirements"],
    summary="Delete a global requirement",
)
async def delete_global_requirement(
    study_id: str,
    req_id: str,
    db: AsyncSession = Depends(get_db),
):
    await _get_study(study_id, db)
    req = await db.get(GlobalRequirement, req_id)
    if req is None or req.study_id != study_id:
        raise _not_found("GlobalRequirement", req_id)
    await db.delete(req)


@app.put(
    "/studies/{study_id}/global-requirements",
    response_model=GlobalRequirementList,
    tags=["Global Requirements"],
    summary="Bulk-replace all global requirements for a study",
)
async def bulk_replace_global_requirements(
    study_id: str,
    body: GlobalRequirementBulkUpdate,
    db: AsyncSession = Depends(get_db),
):
    """
    PUT /studies/{study_id}/global-requirements

    Replaces the entire set of global requirements for this study.
    Existing requirements are deleted and re-created from the payload.
    Use this for the 'Save Requirements' action in the Global Requirements tab.
    """
    await _get_study(study_id, db)

    await db.execute(delete(GlobalRequirement).where(GlobalRequirement.study_id == study_id))

    new_reqs: List[GlobalRequirement] = []
    for item in body.requirements:
        req = GlobalRequirement(
            id=str(uuid.uuid4()),
            study_id=study_id,
            **item.model_dump(),
        )
        db.add(req)
        new_reqs.append(req)

    await _log_action(
        db,
        study_id=study_id,
        action_type="global_reqs_saved",
        target="global_requirements",
        after={
            "count": len(new_reqs),
            "sections": [r.section_type for r in new_reqs],
        },
    )

    await db.flush()
    for req in new_reqs:
        await db.refresh(req)

    return GlobalRequirementList(requirements=new_reqs, total=len(new_reqs))


# ===========================================================================
# Shells  (per-shell AI-generated mock shells, PRD §4.4)
# ===========================================================================

@app.get(
    "/studies/{study_id}/shells",
    response_model=ShellList,
    tags=["Shells"],
    summary="List all shells for a study (React sidebar data)",
)
async def list_shells(
    study_id: str,
    tlf_id: Optional[str] = Query(None, description="Filter by TLF"),
    status_filter: Optional[str] = Query(None, alias="status"),
    db: AsyncSession = Depends(get_db),
):
    """
    GET /studies/{study_id}/shells

    Returns all shells for the React sidebar.
    Optional query params:
      - tlf_id : show shells for a specific TLF only
      - status : draft | in_review | approved
    """
    await _get_study(study_id, db)

    stmt = select(Shell).where(Shell.study_id == study_id)
    if tlf_id:
        stmt = stmt.where(Shell.tlf_id == tlf_id)
    if status_filter:
        stmt = stmt.where(Shell.status == status_filter)

    stmt = stmt.order_by(Shell.created_at)
    result = await db.execute(stmt)
    rows = result.scalars().all()
    return ShellList(shells=list(rows), total=len(rows))


@app.post(
    "/studies/{study_id}/shells",
    response_model=ShellRead,
    status_code=status.HTTP_201_CREATED,
    tags=["Shells"],
    summary="Create a new shell for a TLF",
)
async def create_shell(
    study_id: str,
    body: ShellCreate,
    db: AsyncSession = Depends(get_db),
):
    """
    POST /studies/{study_id}/shells

    Creates a new shell record.  Consults Global Requirements to apply
    section-level defaults (columns, title template) when a TLF is linked.
    In production this would trigger the Biostat Expert → Builder → Reviewer
    AI loop (PRD §4.4).

    Precedence:
      1. Explicit SAP content / user overrides in the request body
      2. Global Requirements defaults for the matching section type
      3. Fallback built-in defaults
    """
    await _get_study(study_id, db)

    applied_columns = body.columns
    applied_title = body.title
    applied_subtitle = body.subtitle
    gr_used: Optional[str] = None  # section_type of matched global requirement

    if body.tlf_id:
        tlf = await _get_tlf(study_id, body.tlf_id, db)

        # Consult global requirements for the TLF's section
        if tlf.section_ref:
            req_result = await db.execute(
                select(GlobalRequirement)
                .where(GlobalRequirement.study_id == study_id)
                .where(GlobalRequirement.section_type == tlf.section_ref)
            )
            matched_req = req_result.scalars().first()

            if matched_req:
                gr_used = matched_req.section_type

                # Apply columns from global requirements only if not explicitly provided
                if not body.columns and matched_req.columns:
                    applied_columns = matched_req.columns

                # Apply title template if title is generic/empty
                if matched_req.title_template and body.title in ("New Shell", "", None):
                    applied_title = (
                        matched_req.title_template
                        .replace("{number}", tlf.number or "")
                        .replace("{title}", tlf.title or "")
                        .replace("{population}", body.population or "")
                    )

                # Apply subtitle template if not explicitly provided
                if matched_req.subtitle_template and not body.subtitle:
                    applied_subtitle = matched_req.subtitle_template

    shell = Shell(
        id=str(uuid.uuid4()),
        study_id=study_id,
        tlf_id=body.tlf_id,
        type=body.type.value,
        title=applied_title,
        subtitle=applied_subtitle,
        population=body.population,
        columns=applied_columns,
        rows=body.rows,
        footnotes=body.footnotes,
        status="draft",
        ai_generation_log={
            "note": "AI generation stub — replace with per-shell loop",
            "global_requirement_applied": gr_used,
        },
    )
    db.add(shell)

    await _log_action(
        db,
        study_id=study_id,
        action_type="shell_created",
        tlf_id=body.tlf_id,
        shell_id=shell.id,
        target="shell",
        after={"title": shell.title, "type": shell.type, "status": shell.status},
    )

    await db.flush()
    await db.refresh(shell)
    return shell


@app.get(
    "/studies/{study_id}/shells/{shell_id}",
    response_model=ShellRead,
    tags=["Shells"],
    summary="Get a single shell",
)
async def get_shell(
    study_id: str,
    shell_id: str,
    db: AsyncSession = Depends(get_db),
):
    return await _get_shell(study_id, shell_id, db)


@app.put(
    "/studies/{study_id}/shells/{shell_id}",
    response_model=ShellRead,
    tags=["Shells"],
    summary="Update a shell (updateActiveShell)",
)
async def update_shell(
    study_id: str,
    shell_id: str,
    body: ShellUpdate,
    db: AsyncSession = Depends(get_db),
):
    """
    PUT /studies/{study_id}/shells/{shell_id}

    Partial update — only provided fields are changed.
    Logs an audit action for every field updated.
    Use to:
      - Edit title / subtitle / population
      - Replace columns / rows / footnotes after user table edits
      - Transition status (draft → in_review → approved)
    """
    shell = await _get_shell(study_id, shell_id, db)

    updates = body.model_dump(exclude_none=True)
    for field, new_value in updates.items():
        old_value = getattr(shell, field, None)
        # Resolve enum → string for storage
        store_value = new_value.value if hasattr(new_value, "value") else new_value
        setattr(shell, field, store_value)

        # Choose a meaningful action_type for the audit log
        action_map = {
            "title": "update_title",
            "subtitle": "update_subtitle",
            "population": "update_population",
            "columns": "update_column",
            "rows": "update_row",
            "footnotes": "update_footnote",
            "status": "update_status",
        }
        action_type = action_map.get(field, "update_row")

        await _log_action(
            db,
            study_id=study_id,
            action_type=action_type,
            tlf_id=shell.tlf_id,
            shell_id=shell_id,
            target=field,
            before={field: old_value},
            after={field: store_value},
        )

    shell.updated_at = datetime.utcnow()
    await db.flush()
    await db.refresh(shell)
    return shell


@app.delete(
    "/studies/{study_id}/shells/{shell_id}",
    status_code=status.HTTP_204_NO_CONTENT,
    tags=["Shells"],
    summary="Delete a shell",
)
async def delete_shell(
    study_id: str,
    shell_id: str,
    db: AsyncSession = Depends(get_db),
):
    shell = await _get_shell(study_id, shell_id, db)
    await _log_action(
        db,
        study_id=study_id,
        action_type="reject_shell",
        tlf_id=shell.tlf_id,
        shell_id=shell_id,
        target="shell",
        before={"title": shell.title, "status": shell.status},
    )
    await db.delete(shell)


# ===========================================================================
# Chat  (AI interaction log, PRD §5.1)
# ===========================================================================

@app.post(
    "/studies/{study_id}/chat",
    response_model=ChatResponse,
    status_code=status.HTTP_201_CREATED,
    tags=["Chat"],
    summary="Send a message and get an AI response stub",
)
async def post_chat(
    study_id: str,
    body: ChatRequest,
    db: AsyncSession = Depends(get_db),
):
    """
    POST /studies/{study_id}/chat

    Stores the user message and a stub AI response.
    In production: forward context + prompt to the per-shell AI loop
    (Biostat Expert → Builder → Reviewer, PRD §4.4), then stream reply.

    Body (either form accepted):
        {
            "tlf_id":    "optional-tlf-uuid",
            "shell_id":  "optional-shell-uuid",
            "prompt":    "Add a footnote referencing SAP Section 12.1"
        }
    or:
        {
            "target":    "shell",
            "target_id": "<shell_id>",
            "prompt":    "Add a footnote referencing SAP Section 12.1"
        }
    """
    await _get_study(study_id, db)

    # Resolve target/target_id routing to concrete FK fields
    resolved_shell_id = body.shell_id
    resolved_tlf_id = body.tlf_id

    if body.target == "shell" and body.target_id:
        resolved_shell_id = body.target_id
    elif body.target == "tlf" and body.target_id:
        resolved_tlf_id = body.target_id

    # Build shell context for AI response generation
    shell_context = ""
    if resolved_shell_id:
        shell_row = await db.get(Shell, resolved_shell_id)
        if shell_row and shell_row.study_id == study_id:
            col_labels = [c.get("label", "") for c in (shell_row.columns or [])]
            row_labels = [r.get("label", "") for r in (shell_row.rows or [])[:10]]
            footnotes = shell_row.footnotes or []
            shell_context = (
                f"\n\nShell context:"
                f"\n  Title: {shell_row.title}"
                f"\n  Type: {shell_row.type}"
                f"\n  Population: {shell_row.population or 'Not specified'}"
                f"\n  Columns: {', '.join(col_labels) if col_labels else 'None'}"
                f"\n  Rows (first 10): {', '.join(row_labels) if row_labels else 'None'}"
                f"\n  Footnotes: {'; '.join(footnotes) if footnotes else 'None'}"
            )

    # Persist user message
    user_msg = Message(
        id=str(uuid.uuid4()),
        study_id=study_id,
        tlf_id=resolved_tlf_id,
        shell_id=resolved_shell_id,
        role="user",
        text=body.prompt,
        timestamp=datetime.utcnow(),
    )
    db.add(user_msg)

    # Log audit event for chat interaction
    await _log_action(
        db,
        study_id=study_id,
        action_type="chat_sent",
        shell_id=resolved_shell_id,
        tlf_id=resolved_tlf_id,
        target=body.target or "study",
        after={
            "prompt": body.prompt,
            "shell_id": resolved_shell_id,
            "target": body.target,
        },
    )

    # --- AI stub -------------------------------------------------------
    # Real implementation: retrieve top-k chunks from parsed documents,
    # run multi-role AI loop, return structured shell update + explanation.
    ai_text = (
        f'[AI STUB] Received: "{body.prompt}". '
        "In production this triggers the Biostat Expert -> Builder -> "
        f"Reviewer loop and returns structured shell updates.{shell_context}"
    )
    ai_msg = Message(
        id=str(uuid.uuid4()),
        study_id=study_id,
        tlf_id=resolved_tlf_id,
        shell_id=resolved_shell_id,
        role="ai",
        text=ai_text,
        timestamp=datetime.utcnow(),
        extra_metadata={
            "model": "stub",
            "shell_id": resolved_shell_id,
            "retrieved_chunks": [],
            "confidence": None,
        },
    )
    db.add(ai_msg)
    # -------------------------------------------------------------------

    await db.flush()
    await db.refresh(user_msg)
    await db.refresh(ai_msg)

    return ChatResponse(
        user_message=MessageRead.model_validate(user_msg),
        ai_message=MessageRead.model_validate(ai_msg),
    )


@app.get(
    "/studies/{study_id}/chat",
    response_model=MessageList,
    tags=["Chat"],
    summary="Get chat history for a study (optionally filtered by TLF/shell)",
)
async def get_chat_history(
    study_id: str,
    tlf_id: Optional[str] = Query(None),
    shell_id: Optional[str] = Query(None),
    skip: int = Query(0, ge=0),
    limit: int = Query(100, ge=1, le=500),
    db: AsyncSession = Depends(get_db),
):
    await _get_study(study_id, db)

    stmt = (
        select(Message)
        .where(Message.study_id == study_id)
        .order_by(Message.timestamp)
    )
    if tlf_id:
        stmt = stmt.where(Message.tlf_id == tlf_id)
    if shell_id:
        stmt = stmt.where(Message.shell_id == shell_id)

    stmt = stmt.offset(skip).limit(limit)
    result = await db.execute(stmt)
    rows = result.scalars().all()
    return MessageList(messages=list(rows), total=len(rows))


@app.get(
    "/studies/{study_id}/shells/{shell_id}/messages",
    response_model=MessageList,
    tags=["Chat"],
    summary="Get chat history for a specific shell",
)
async def get_shell_messages(
    study_id: str,
    shell_id: str,
    skip: int = Query(0, ge=0),
    limit: int = Query(100, ge=1, le=500),
    db: AsyncSession = Depends(get_db),
):
    """
    GET /studies/{study_id}/shells/{shell_id}/messages

    Returns ordered chat history (user + AI messages) for a specific shell.
    Used by the frontend to load per-shell persistent conversation history.
    """
    await _get_shell(study_id, shell_id, db)

    stmt = (
        select(Message)
        .where(Message.study_id == study_id)
        .where(Message.shell_id == shell_id)
        .order_by(Message.timestamp)
        .offset(skip)
        .limit(limit)
    )
    result = await db.execute(stmt)
    rows = result.scalars().all()
    return MessageList(messages=list(rows), total=len(rows))


# ===========================================================================
# Audit Events  (enriched audit trail for UI history panels)
# ===========================================================================

@app.get(
    "/studies/{study_id}/audit-events",
    response_model=AuditEventList,
    tags=["Audit"],
    summary="Get enriched study-level audit events for the UI history panel",
)
async def get_audit_events(
    study_id: str,
    limit: int = Query(50, ge=1, le=500),
    skip: int = Query(0, ge=0),
    db: AsyncSession = Depends(get_db),
):
    """
    GET /studies/{study_id}/audit-events

    Returns study-level audit events newest-first with human-readable summaries,
    entity type labels, and source badges.  Intended for the UI audit history panel.
    """
    await _get_study(study_id, db)

    # Total count
    count_result = await db.execute(
        select(func.count(Action.id)).where(Action.study_id == study_id)
    )
    total = count_result.scalar_one()

    stmt = (
        select(Action)
        .where(Action.study_id == study_id)
        .order_by(Action.timestamp.desc())
        .offset(skip)
        .limit(limit)
    )
    result = await db.execute(stmt)
    rows = result.scalars().all()
    events = [_action_to_event(a) for a in rows]
    return AuditEventList(events=events, total=total)


@app.get(
    "/studies/{study_id}/shells/{shell_id}/audit-events",
    response_model=AuditEventList,
    tags=["Audit"],
    summary="Get enriched shell-level audit events for the shell history panel",
)
async def get_shell_audit_events(
    study_id: str,
    shell_id: str,
    limit: int = Query(50, ge=1, le=500),
    skip: int = Query(0, ge=0),
    db: AsyncSession = Depends(get_db),
):
    """
    GET /studies/{study_id}/shells/{shell_id}/audit-events

    Returns shell-specific audit events newest-first.
    Includes shell edits, status changes, chat interactions, and AI suggestions.
    """
    await _get_shell(study_id, shell_id, db)

    count_result = await db.execute(
        select(func.count(Action.id))
        .where(Action.study_id == study_id)
        .where(Action.shell_id == shell_id)
    )
    total = count_result.scalar_one()

    stmt = (
        select(Action)
        .where(Action.study_id == study_id)
        .where(Action.shell_id == shell_id)
        .order_by(Action.timestamp.desc())
        .offset(skip)
        .limit(limit)
    )
    result = await db.execute(stmt)
    rows = result.scalars().all()
    events = [_action_to_event(a) for a in rows]
    return AuditEventList(events=events, total=total)


# ===========================================================================
# Audit trail  (GxP-style action log, PRD §5.1)
# ===========================================================================

@app.get(
    "/studies/{study_id}/audit",
    response_model=ActionList,
    tags=["Audit"],
    summary="Get the audit trail for a study",
)
async def get_audit_trail(
    study_id: str,
    tlf_id: Optional[str] = Query(None),
    shell_id: Optional[str] = Query(None),
    skip: int = Query(0, ge=0),
    limit: int = Query(200, ge=1, le=1000),
    db: AsyncSession = Depends(get_db),
):
    """
    GET /studies/{study_id}/audit

    Returns every logged action (user edits + AI suggestions).
    Exportable to CSV for the compliance audit log (PRD §3.1).
    """
    await _get_study(study_id, db)

    stmt = (
        select(Action)
        .where(Action.study_id == study_id)
        .order_by(Action.timestamp.desc())
    )
    if tlf_id:
        stmt = stmt.where(Action.tlf_id == tlf_id)
    if shell_id:
        stmt = stmt.where(Action.shell_id == shell_id)

    stmt = stmt.offset(skip).limit(limit)
    result = await db.execute(stmt)
    rows = result.scalars().all()
    return ActionList(actions=list(rows), total=len(rows))


@app.post(
    "/studies/{study_id}/audit",
    response_model=ActionRead,
    status_code=status.HTTP_201_CREATED,
    tags=["Audit"],
    summary="Log a manual audit action",
)
async def log_audit_action(
    study_id: str,
    body: ActionCreate,
    db: AsyncSession = Depends(get_db),
):
    """
    POST /studies/{study_id}/audit

    Explicitly log a user or AI action (e.g. direct table edits from React UI).
    The backend also logs actions automatically on shell/TLF mutations.
    """
    await _get_study(study_id, db)

    action = await _log_action(
        db,
        study_id=study_id,
        action_type=body.type.value,
        tlf_id=body.tlf_id,
        shell_id=body.shell_id,
        target=body.target,
        before=body.before_state,
        after=body.after_state,
        actor=body.actor,
    )
    await db.flush()
    await db.refresh(action)
    return action


# ===========================================================================
# Export endpoints  (PRD §3.1 — study outputs for handoff and regulatory)
# ===========================================================================

def _shell_to_export_dict(shell: Shell) -> dict:
    """Convert a Shell ORM object to a clean export dict."""
    return {
        "id": shell.id,
        "tlf_id": shell.tlf_id,
        "type": shell.type,
        "title": shell.title,
        "subtitle": shell.subtitle,
        "population": shell.population,
        "status": shell.status,
        "columns": shell.columns or [],
        "rows": shell.rows or [],
        "footnotes": shell.footnotes or [],
        "created_at": shell.created_at.isoformat() if shell.created_at else None,
        "updated_at": shell.updated_at.isoformat() if shell.updated_at else None,
    }


def _tlf_to_export_dict(tlf: TLF) -> dict:
    return {
        "id": tlf.id,
        "number": tlf.number,
        "title": tlf.title,
        "type": tlf.type,
        "section_ref": tlf.section_ref,
        "status": tlf.status,
        "order_index": tlf.order_index,
    }


def _req_to_export_dict(req: GlobalRequirement) -> dict:
    return {
        "id": req.id,
        "section_type": req.section_type,
        "number_pattern": req.number_pattern,
        "title_template": req.title_template,
        "subtitle_template": req.subtitle_template,
        "columns": req.columns or [],
    }


def _action_to_export_dict(action: Action) -> dict:
    return {
        "id": action.id,
        "timestamp": action.timestamp.isoformat() if action.timestamp else None,
        "type": action.type,
        "entity_type": action.entity_type,
        "target": action.target,
        "summary": action.summary,
        "actor": action.actor,
        "tlf_id": action.tlf_id,
        "shell_id": action.shell_id,
    }


@app.get(
    "/studies/{study_id}/export/json",
    tags=["Export"],
    summary="Export full study as structured JSON",
)
async def export_study_json(
    study_id: str,
    db: AsyncSession = Depends(get_db),
):
    """
    GET /studies/{study_id}/export/json

    Returns a complete JSON export of the study: metadata, TLF list,
    global requirements, shells, and audit events.
    Suitable for archival, programmer handoff, and regulatory review.
    """
    study = await _get_study(study_id, db)

    tlfs_result = await db.execute(
        select(TLF).where(TLF.study_id == study_id).order_by(TLF.order_index)
    )
    tlfs = tlfs_result.scalars().all()

    reqs_result = await db.execute(
        select(GlobalRequirement).where(GlobalRequirement.study_id == study_id)
    )
    reqs = reqs_result.scalars().all()

    shells_result = await db.execute(
        select(Shell).where(Shell.study_id == study_id).order_by(Shell.created_at)
    )
    shells = shells_result.scalars().all()

    actions_result = await db.execute(
        select(Action)
        .where(Action.study_id == study_id)
        .order_by(Action.timestamp.desc())
        .limit(500)
    )
    actions = actions_result.scalars().all()

    payload = {
        "export_version": "1.0",
        "export_timestamp": datetime.utcnow().isoformat() + "Z",
        "study": {
            "id": study.id,
            "name": study.name,
            "created_at": study.created_at.isoformat() if study.created_at else None,
            "updated_at": study.updated_at.isoformat() if study.updated_at else None,
        },
        "tlf_list": [_tlf_to_export_dict(t) for t in tlfs],
        "global_requirements": [_req_to_export_dict(r) for r in reqs],
        "shells": [_shell_to_export_dict(s) for s in shells],
        "audit_events": [_action_to_export_dict(a) for a in actions],
    }

    safe_name = study.name.replace(" ", "_").replace("/", "_")[:40]
    filename = f"study_{safe_name}_export.json"

    json_bytes = json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8")

    return StreamingResponse(
        io.BytesIO(json_bytes),
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get(
    "/studies/{study_id}/export/tlf-list.csv",
    tags=["Export"],
    summary="Export TLF list as CSV",
)
async def export_tlf_list_csv(
    study_id: str,
    db: AsyncSession = Depends(get_db),
):
    """
    GET /studies/{study_id}/export/tlf-list.csv

    Returns TLF list as CSV with stable columns:
    number, title, type, section, status
    """
    study = await _get_study(study_id, db)

    result = await db.execute(
        select(TLF).where(TLF.study_id == study_id).order_by(TLF.order_index)
    )
    tlfs = result.scalars().all()

    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["number", "title", "type", "section", "status"])
    for t in tlfs:
        writer.writerow([
            t.number or "",
            t.title or "",
            t.type or "",
            t.section_ref or "",
            t.status or "",
        ])

    safe_name = study.name.replace(" ", "_").replace("/", "_")[:40]
    filename = f"study_{safe_name}_tlf_list.csv"

    return StreamingResponse(
        io.BytesIO(buf.getvalue().encode("utf-8")),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get(
    "/studies/{study_id}/export/shells.json",
    tags=["Export"],
    summary="Export all shells as JSON for programmer consumption",
)
async def export_shells_json(
    study_id: str,
    db: AsyncSession = Depends(get_db),
):
    """
    GET /studies/{study_id}/export/shells.json

    Returns all shells with complete structure (title, subtitle, columns,
    rows, footnotes). Suitable for SAS/R programmer handoff.
    """
    study = await _get_study(study_id, db)

    result = await db.execute(
        select(Shell).where(Shell.study_id == study_id).order_by(Shell.created_at)
    )
    shells = result.scalars().all()

    payload = {
        "export_version": "1.0",
        "export_timestamp": datetime.utcnow().isoformat() + "Z",
        "study_id": study_id,
        "study_name": study.name,
        "shells": [_shell_to_export_dict(s) for s in shells],
    }

    safe_name = study.name.replace(" ", "_").replace("/", "_")[:40]
    filename = f"study_{safe_name}_shells.json"

    json_bytes = json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8")

    return StreamingResponse(
        io.BytesIO(json_bytes),
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get(
    "/studies/{study_id}/export/audit.csv",
    tags=["Export"],
    summary="Export audit trail as CSV",
)
async def export_audit_csv(
    study_id: str,
    db: AsyncSession = Depends(get_db),
):
    """
    GET /studies/{study_id}/export/audit.csv

    Returns audit events as CSV. Suitable for regulatory compliance review.
    Columns: timestamp, entity_type, entity_id, action, summary, actor, target
    """
    study = await _get_study(study_id, db)

    result = await db.execute(
        select(Action)
        .where(Action.study_id == study_id)
        .order_by(Action.timestamp.desc())
        .limit(2000)
    )
    actions = result.scalars().all()

    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow([
        "timestamp", "entity_type", "entity_id", "action", "summary", "actor", "target"
    ])
    for a in actions:
        entity_id = a.shell_id or a.tlf_id or a.study_id or ""
        writer.writerow([
            a.timestamp.isoformat() if a.timestamp else "",
            a.entity_type or "",
            entity_id,
            a.type or "",
            a.summary or "",
            a.actor or "",
            a.target or "",
        ])

    safe_name = study.name.replace(" ", "_").replace("/", "_")[:40]
    filename = f"study_{safe_name}_audit.csv"

    return StreamingResponse(
        io.BytesIO(buf.getvalue().encode("utf-8")),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get(
    "/studies/{study_id}/export/shells/{shell_id}.docx",
    tags=["Export"],
    summary="Export a single shell as a Word document",
)
async def export_shell_docx(
    study_id: str,
    shell_id: str,
    db: AsyncSession = Depends(get_db),
):
    """
    GET /studies/{study_id}/export/shells/{shell_id}.docx

    Returns a single shell as a Word .docx file with:
    - Title at top (bold, centered)
    - Subtitle / population line
    - Table grid with column headers and row stubs
    - Footnotes below

    Suitable for programmer handoff and regulatory review.
    """
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise HTTPException(
            status_code=501,
            detail="python-docx not installed. Run: pip install python-docx",
        )

    shell = await _get_shell(study_id, shell_id, db)

    doc = DocxDocument()

    # --- Page margins ---
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Title ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(shell.title or "Untitled Shell")
    title_run.bold = True
    title_run.font.size = Pt(12)

    # --- Subtitle ---
    if shell.subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(shell.subtitle)
        sub_run.font.size = Pt(11)

    # --- Population line ---
    if shell.population:
        pop_para = doc.add_paragraph()
        pop_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pop_run = pop_para.add_run(f"Population: {shell.population}")
        pop_run.font.size = Pt(10)
        pop_run.italic = True

    doc.add_paragraph()  # spacer

    # --- Table ---
    columns = shell.columns or []
    rows = shell.rows or []

    if not columns:
        # Fallback: stub table with a single column
        columns = [{"label": "Row Stub"}]

    num_cols = len(columns) + 1  # +1 for row label column
    tbl = doc.add_table(rows=1, cols=num_cols)
    tbl.style = "Table Grid"

    # Header row
    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].text = "Row Stub"
    hdr_cells[0].paragraphs[0].runs[0].bold = True
    for i, col in enumerate(columns):
        cell = hdr_cells[i + 1]
        cell.text = col.get("label", f"Col {i+1}")
        cell.paragraphs[0].runs[0].bold = True

    # Data rows
    for row in rows:
        cells = tbl.add_row().cells
        label = row.get("label", "")
        indent = row.get("indent", 0)
        is_header = row.get("is_header", False)

        # Indent with spaces
        cells[0].text = ("  " * indent) + label
        if is_header:
            for run in cells[0].paragraphs[0].runs:
                run.bold = True

        # Fill data columns with placeholder
        for i in range(len(columns)):
            cells[i + 1].text = "X"

    doc.add_paragraph()  # spacer

    # --- Footnotes ---
    footnotes = shell.footnotes or []
    if footnotes:
        fn_para = doc.add_paragraph()
        fn_run = fn_para.add_run("Footnotes:")
        fn_run.bold = True
        fn_run.font.size = Pt(9)
        for fn in footnotes:
            fn_p = doc.add_paragraph(style="List Bullet")
            fn_run2 = fn_p.add_run(str(fn))
            fn_run2.font.size = Pt(9)

    # --- Serialize to bytes ---
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = (shell.title or "shell").replace(" ", "_").replace("/", "_")[:50]
    filename = f"shell_{safe_title}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get(
    "/studies/{study_id}/export/package.zip",
    tags=["Export"],
    summary="Export complete study package as ZIP bundle",
)
async def export_study_package_zip(
    study_id: str,
    db: AsyncSession = Depends(get_db),
):
    """
    GET /studies/{study_id}/export/package.zip

    Returns a ZIP file containing all study artifacts:
    - study.json  (complete study export)
    - tlf_list.csv
    - global_requirements.json
    - shells.json
    - audit.csv
    """
    study = await _get_study(study_id, db)

    # Load all data
    tlfs_result = await db.execute(
        select(TLF).where(TLF.study_id == study_id).order_by(TLF.order_index)
    )
    tlfs = tlfs_result.scalars().all()

    reqs_result = await db.execute(
        select(GlobalRequirement).where(GlobalRequirement.study_id == study_id)
    )
    reqs = reqs_result.scalars().all()

    shells_result = await db.execute(
        select(Shell).where(Shell.study_id == study_id).order_by(Shell.created_at)
    )
    shells = shells_result.scalars().all()

    actions_result = await db.execute(
        select(Action)
        .where(Action.study_id == study_id)
        .order_by(Action.timestamp.desc())
        .limit(2000)
    )
    actions = actions_result.scalars().all()

    export_ts = datetime.utcnow().isoformat() + "Z"

    # --- study.json ---
    study_json = json.dumps({
        "export_version": "1.0",
        "export_timestamp": export_ts,
        "study": {
            "id": study.id,
            "name": study.name,
            "created_at": study.created_at.isoformat() if study.created_at else None,
            "updated_at": study.updated_at.isoformat() if study.updated_at else None,
        },
        "tlf_list": [_tlf_to_export_dict(t) for t in tlfs],
        "global_requirements": [_req_to_export_dict(r) for r in reqs],
        "shells": [_shell_to_export_dict(s) for s in shells],
        "audit_events": [_action_to_export_dict(a) for a in actions],
    }, indent=2, ensure_ascii=False).encode("utf-8")

    # --- tlf_list.csv ---
    tlf_buf = io.StringIO()
    tlf_writer = csv.writer(tlf_buf)
    tlf_writer.writerow(["number", "title", "type", "section", "status"])
    for t in tlfs:
        tlf_writer.writerow([t.number or "", t.title or "", t.type or "",
                              t.section_ref or "", t.status or ""])
    tlf_csv = tlf_buf.getvalue().encode("utf-8")

    # --- global_requirements.json ---
    reqs_json = json.dumps(
        [_req_to_export_dict(r) for r in reqs], indent=2, ensure_ascii=False
    ).encode("utf-8")

    # --- shells.json ---
    shells_json = json.dumps({
        "export_version": "1.0",
        "export_timestamp": export_ts,
        "study_id": study_id,
        "study_name": study.name,
        "shells": [_shell_to_export_dict(s) for s in shells],
    }, indent=2, ensure_ascii=False).encode("utf-8")

    # --- audit.csv ---
    audit_buf = io.StringIO()
    audit_writer = csv.writer(audit_buf)
    audit_writer.writerow([
        "timestamp", "entity_type", "entity_id", "action", "summary", "actor", "target"
    ])
    for a in actions:
        entity_id = a.shell_id or a.tlf_id or a.study_id or ""
        audit_writer.writerow([
            a.timestamp.isoformat() if a.timestamp else "",
            a.entity_type or "", entity_id, a.type or "",
            a.summary or "", a.actor or "", a.target or "",
        ])
    audit_csv = audit_buf.getvalue().encode("utf-8")

    # --- Build ZIP ---
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("study.json", study_json)
        zf.writestr("tlf_list.csv", tlf_csv)
        zf.writestr("global_requirements.json", reqs_json)
        zf.writestr("shells.json", shells_json)
        zf.writestr("audit.csv", audit_csv)

    zip_buf.seek(0)

    safe_name = study.name.replace(" ", "_").replace("/", "_")[:40]
    filename = f"study_{safe_name}_package.zip"

    return StreamingResponse(
        zip_buf,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
