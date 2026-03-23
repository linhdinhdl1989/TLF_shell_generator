"""Shared pytest fixtures for backend tests."""
import io
import pytest
import openpyxl
from docx import Document as DocxDocument


SAP_SAMPLE_TEXT = """\
Statistical Analysis Plan
Study: XYZ-101

1. Introduction
This SAP describes the statistical methods for Study XYZ-101.

5. Analysis Populations
The primary analysis will use the Safety Population.
All randomized subjects who received at least one dose are included.

12. List of Tables, Listings, and Figures

Table 14.1.1 Summary of Demographic and Baseline Characteristics
This table summarizes demographics. Summaries will be presented by Treatment Group.
Subjects are analyzed using the Safety Population.

Table 14.1.2 Summary of Prior and Concomitant Medications

Table 14.2.1 Primary Efficacy Endpoint Analysis
The primary endpoint analysis. Intent-to-Treat Population.

Listing 16.1.1 Subject Disposition Listing

Table 14.3.1 Overview of Adverse Events (Safety Population)
Adverse events summaries will be presented by System Organ Class and Preferred Term.

Table 14.3.2 Treatment-Emergent Adverse Events by System Organ Class
"""


@pytest.fixture
def sap_text():
    """Raw SAP text for extraction tests."""
    return SAP_SAMPLE_TEXT


@pytest.fixture
def minimal_docx_bytes():
    """In-memory .docx with known content."""
    doc = DocxDocument()
    doc.add_heading("Statistical Analysis Plan", level=1)
    doc.add_paragraph("5. Analysis Populations")
    doc.add_paragraph(
        "Table 14.1.1 Summary of Demographic and Baseline Characteristics"
    )
    doc.add_paragraph(
        "Table 14.2.1 Primary Efficacy Endpoint Analysis (Intent-to-Treat Population)"
    )
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@pytest.fixture
def minimal_xlsx_bytes():
    """In-memory .xlsx with known content."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "TLF List"
    ws.append(["Number", "Title", "Section"])
    ws.append(["14.1.1", "Demographics", "demographics"])
    ws.append(["14.2.1", "Primary Efficacy", "efficacy"])
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()
